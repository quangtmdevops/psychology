from docx import Document
import re
from sqlalchemy.orm import Session
from app.models.models import Test, Entity, Group, TestAnswer, Option
from typing import List, Dict, Any

class QuestionService:
    @staticmethod
    def read_questions_from_docx(file_path: str) -> List[Dict[str, Any]]:
        doc = Document(file_path)
        questions = []
        print(f"Đang đọc file: {file_path}")
        print(f"Số bảng trong file: {len(doc.tables)}")
        code_col_idx = None
        is_mdq = 'SÀNG-LỌC-RỐI-LOẠN-CẢM-XÚC-LƯỠNG-CỰC' in file_path or 'MDQ' in file_path.upper()
        mdq_yesno_numbers = {"1", "2"} | {f"1.{i}" for i in range(1, 14)}
        # Đọc từ bảng đầu tiên
        if len(doc.tables) > 0:
            table = doc.tables[0]
            print(f"Số hàng trong bảng: {len(table.rows)}")
            # In ra header để kiểm tra
            if len(table.rows) > 0:
                header = [cell.text.strip() for cell in table.rows[0].cells]
                print(f"Header của bảng: {header}")
                # Tìm vị trí cột "Mã" nếu có
                for idx, col in enumerate(header):
                    if col.lower() == "mã":
                        code_col_idx = idx
                        break
            # Bỏ qua hàng header
            for row_idx, row in enumerate(table.rows[1:], 1):
                cells = row.cells
                print(f"\nĐang đọc hàng {row_idx}:")
                for cell_idx, cell in enumerate(cells):
                    print(f"Cột {cell_idx}: {cell.text.strip()}")
                if len(cells) >= 2:
                    statement_text = cells[1].text.strip()
                    code_value = None
                    if code_col_idx is not None and code_col_idx < len(cells):
                        code_value = cells[code_col_idx].text.strip() or None
                    # Lấy số thứ tự câu hỏi (thường ở cột 0)
                    question_number = cells[0].text.strip()
                    if statement_text:
                        # Xử lý riêng cho file MDQ
                        if is_mdq:
                            if question_number in mdq_yesno_numbers:
                                options = [
                                    {'level': 0, 'content': 'Không'},
                                    {'level': 1, 'content': 'Có'}
                                ]
                            elif 'nghiêm trọng ở mức độ nào' in statement_text.lower():
                                options = [
                                    {'level': 1, 'content': '1. Không nghiêm trọng'},
                                    {'level': 2, 'content': '2. Vừa phải'},
                                    {'level': 3, 'content': '3. Khá nghiêm trọng'},
                                    {'level': 4, 'content': '4. Rất nghiêm trọng'}
                                ]
                            else:
                                options = [
                                    {'level': 0, 'content': 'Không'},
                                    {'level': 1, 'content': 'Có'}
                                ]
                        else:
                            options = [
                                {'level': 0, 'content': 'Không đúng với tôi chút nào cả'},
                                {'level': 1, 'content': 'Đúng với tôi phần nào, hoặc thỉnh thoảng mới đúng'},
                                {'level': 2, 'content': 'Đúng với tôi phần nhiều, hoặc phần lớn thời gian là đúng'},
                                {'level': 3, 'content': 'Hoàn toàn đúng với tôi, hoặc hầu hết thời gian là đúng'}
                            ]
                        questions.append({
                            'content': statement_text,
                            'code': code_value,
                            'options': options
                        })
                        print(f"Đã thêm câu tự sự: {statement_text}, mã: {code_value}, options: {options}")
                    else:
                        print(f"Bỏ qua hàng {row_idx} vì không có nội dung")
        print(f"\nTổng số câu tự sự đọc được: {len(questions)}")
        return questions

    @staticmethod
    def import_questions_to_db(file_path: str, group_name: str, db: Session) -> None:
        try:
            # Đọc câu hỏi từ file
            questions = QuestionService.read_questions_from_docx(file_path)
            # Lấy hoặc tạo group
            group = db.query(Group).filter_by(name=group_name).first()
            if not group:
                group = Group(name=group_name)
                db.add(group)
                db.flush()
            for q in questions:
                existing_test = db.query(Test).filter_by(content=q['content'], group_id=group.id).first()
                if existing_test:
                    continue
                test = Test(content=q['content'], group_id=group.id, code=q.get('code'))
                db.add(test)
                db.flush()
                for opt in q['options']:
                    existing_option = db.query(Option).filter_by(test_id=test.id, content=opt['content'], level=opt['level']).first()
                    if existing_option:
                        continue
                    option = Option(test_id=test.id, content=opt['content'], level=opt['level'])
                    db.add(option)
            db.commit()
        except Exception as e:
            db.rollback()
            print(f"Lỗi khi import vào database: {str(e)}")
            raise

    @staticmethod
    def import_all_questions(db: Session) -> None:
        # Đường dẫn tới các file câu hỏi và tên nhóm tương ứng
        file_paths = {
            'DASS': 'app/data/questions/TRẮC-NGHIỆM-LO-ÂU-TRẦM-CẢM.-STRESS-DASS.docx',
            'RADS': 'app/data/questions/TRẮC-NGHIỆM-ĐÁNH-GIÁ-TRẦM-CẢM-THANH-THIẾU-NIÊN-RADS.docx',
            'MDQ': 'app/data/questions/TRẮC-NGHIỆM-SÀNG-LỌC-RỐI-LOẠN-CẢM-XÚC-LƯỠNG-CỰC.docx'
        }
        
        for test_type, file_path in file_paths.items():
            print(f"\nĐang import file {file_path} vào nhóm {test_type}...")
            QuestionService.import_questions_to_db(file_path, test_type, db) 